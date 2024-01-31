# ИДЗ Помощник 1.0
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
import shutil

# Запрос на ввод данных
IDZ_num = input("Введите номер ИДЗ: ")
IDZ_name = input("Введите полной название ИДЗ, с 'ИДЗ №': ")
img_count = int(input("Введите количество картинок: "))
img_path = input("Введите полный путь до папки, где лежат картинки: ")
IDZ_path = input("Введите полный путь до файла, в который вы хотите сохранить файл: ")

# Создание файла

document = Document()

# Создание и корректировка положения первого параграфа

paragraph1 = document.add_paragraph()

run1_1 = paragraph1.add_run("Францев Д.В") # Добавление текста к параграфу
run1_2 = paragraph1.add_run("\nБПИ-23-2\n")

# Добавляем выравнивание

p1_fmt = paragraph1.paragraph_format
p1_fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Создание и корректировка второго параграфа

paragraph2 = document.add_paragraph()

run2_1 = paragraph2.add_run(IDZ_name + "\n" * 2)
run2_2 = paragraph2.add_run("Вариант №21" "\n")

p2_fmt = paragraph2.paragraph_format
p2_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER

#Добавление фотографий

img_path = img_path.split('\\')
img_path = "/".join(img_path)

for i in range(1, img_count+1):
    document.add_picture(img_path + f"/{i}.jpg", width=docx.shared.Cm(18))

# Сохранение файла

document.save("exe.docx")

# Перенос файла в нужную папку

IDZ_path = IDZ_path.split('\\')
IDZ_path = "/".join(IDZ_path)

shutil.move("exe.docx", IDZ_path)

os.rename(IDZ_path + "/exe.docx", IDZ_path + f"/ИДЗ №{IDZ_num}.docx")

  